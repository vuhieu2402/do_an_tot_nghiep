import os
import cv2
import numpy as np
import pytesseract
from docx.oxml import parse_xml
from pytesseract import Output

from process_img.utils import noise_removal, enhance_text
from PIL import Image
from paddleocr import PaddleOCR, draw_ocr
from vietocr.tool.predictor import Predictor
from vietocr.tool.config import Cfg


# Khởi tạo PaddleOCR
ocr = PaddleOCR()

# Khởi tạo VietOCR
config = Cfg.load_config_from_name('vgg_transformer')
# config['weights'] = 'https://vocr.vn/data/vietocr/vgg_transformer.pth'  # Thay bằng đường dẫn đến mô hình
config['weights'] = 'https://vocr.vn/data/vietocr/vgg_transformer.pth'
config['device'] = 'cpu'  # Chạy trên CPU
vietocr = Predictor(config)


def ocr_with_vietocr(image):
    # Chuyển đổi ảnh numpy sang PIL Image nếu cần
    if isinstance(image, np.ndarray):
        image = Image.fromarray(image)

    # Chuyển đổi ảnh sang hệ màu RGB
    image_rgb = image.convert("RGB")

    result = ocr.ocr(np.array(image_rgb))

    if not result or not result[0]:
        return ''

    text_from_image = ''
    for line in result[0]:
        bounding_box = line[0]
        text = line[1][0]
        confidence = line[1][1]

        # Cắt phần ảnh tương ứng với bounding box
        cropped_image = image_rgb.crop((
            int(bounding_box[0][0]),  # x_min
            int(bounding_box[0][1]),  # y_min
            int(bounding_box[2][0]),  # x_max
            int(bounding_box[2][1])  # y_max
        ))

        # Dự đoán từ bằng VietOCR
        predicted_text = vietocr.predict(cropped_image)
        text_from_image += predicted_text + '\n'

    return text_from_image.strip()


def process_uploaded_image(uploaded_files):
    result_items = []  # Tạo danh sách chứa kết quả từ tất cả các ảnh
    doc = Document()  # Tạo file Word mới để lưu kết quả

    for uploaded_file in uploaded_files:
        uploaded_file.seek(0)
        filename = uploaded_file.name

        image_bytes = uploaded_file.read()
        img = cv2.imdecode(np.frombuffer(image_bytes, np.uint8), cv2.IMREAD_GRAYSCALE)

        # Bước 1: Loại bỏ nhiễu và làm nổi bật văn bản

        img = enhance_text(img)


        _, binary_img = cv2.threshold(img, 150, 255, cv2.THRESH_BINARY_INV)

        horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
        vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))

        horizontal_lines = cv2.morphologyEx(binary_img, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
        vertical_lines = cv2.morphologyEx(binary_img, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
        table_area = cv2.add(horizontal_lines, vertical_lines)

        contours, _ = cv2.findContours(table_area, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
        bounding_boxes = [cv2.boundingRect(c) for c in contours]


        if bounding_boxes:
            max_contour = max(bounding_boxes, key=lambda b: b[2] * b[3])
            bounding_boxes = [b for b in bounding_boxes if b != max_contour]

            img_text_only = binary_img.copy()
            for x, y, w, h in bounding_boxes:
                cv2.rectangle(img_text_only, (x, y), (x + w, y + h), (0, 0, 0), -1)

            text_from_image = ocr_with_vietocr(img_text_only)
            result_items.append((float('inf'), text_from_image))
            doc.add_paragraph(text_from_image)

            table_data = []
            for x, y, w, h in bounding_boxes:
                if w > 50 and h > 20:
                    cell_img = img[y:y + h, x:x + w]
                    cell_text = ocr_with_vietocr(cell_img)
                    table_data.append((x, y, cell_text))

            if table_data:
                sorted_table_data = sorted(table_data, key=lambda t: (t[1], t[0]))
                rows = len(set([y for _, y, _ in sorted_table_data]))
                cols = len(set([x for x, _, _ in sorted_table_data]))
                table = doc.add_table(rows=rows, cols=cols)

                row_idx, col_idx = 0, 0
                prev_y = sorted_table_data[0][1]
                for x, y, text in sorted_table_data:
                    if y != prev_y:
                        row_idx += 1
                        col_idx = 0
                        prev_y = y

                    if row_idx < rows and col_idx < cols:
                        cell = table.cell(row_idx, col_idx)
                        cell.text = text
                        cell._element.get_or_add_tcPr().append(parse_xml(
                            r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:left w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/>'
                            r'<w:top w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/>'
                            r'</w:tcBorders>'
                        ))
                        col_idx += 1

        else:
            text_from_image = ocr_with_vietocr(binary_img)
            result_items.append((float('inf'), text_from_image))
            doc.add_paragraph(text_from_image)

    downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
    output_file_path = os.path.join(downloads_folder, 'ocr_result.docx')
    doc.save(output_file_path)

    # Chuyển đổi file .docx thành HTML
    html_content = convert_docx_to_html(output_file_path)

    return html_content  # Trả về nội dung HTML


from docx import Document

def convert_docx_to_html(doc_path):
    doc = Document(doc_path)
    html_content = "<div>"

    # Lặp qua từng đoạn văn bản trong file .docx
    for para in doc.paragraphs:
        html_content += f"<p>{para.text}</p>"

    # Lặp qua từng bảng và tạo cấu trúc HTML cho bảng
    for table in doc.tables:
        html_content += "<table border='1' style='border-collapse: collapse; width: 100%;'>"
        for row in table.rows:
            html_content += "<tr>"
            for cell in row.cells:
                html_content += f"<td>{cell.text}</td>"
            html_content += "</tr>"
        html_content += "</table>"

    html_content += "</div>"
    return html_content
