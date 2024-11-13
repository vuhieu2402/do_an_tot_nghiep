import os
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .forms import PDFUploadForm, DOCXUploadForm
from pdf2docx import Converter
import tempfile
import aspose.words as aw
from googletrans import Translator
from docx import Document
import re



def convert_and_translate_pdf_to_docx_view(request):
    if request.method == 'POST':
        form = PDFUploadForm(request.POST, request.FILES)
        if form.is_valid():
            pdf_file = request.FILES['pdf_file']
            original_filename = pdf_file.name.rsplit('.', 1)[0]  # Original filename without extension
            source_language = form.cleaned_data['source_language']
            target_language = form.cleaned_data['target_language']

            # Save the PDF file to a temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                for chunk in pdf_file.chunks():
                    temp_pdf.write(chunk)
                pdf_path = temp_pdf.name

            # Path for saving the DOCX file
            docx_path = os.path.join(tempfile.gettempdir(), f"{original_filename}.docx")

            try:
                # Convert PDF to DOCX using pdf2docx
                cv = Converter(pdf_path)
                cv.convert(docx_path, start=0, end=None)
                cv.close()

                # Read the content from the DOCX file (if needed for other purposes like translation)
                doc = Document(docx_path)

                # Translate the text using Google Translate API (if needed)
                translated_text = ""
                translator = Translator()

                # Translate and update the text in the DOCX
                for para in doc.paragraphs:
                    if para.text.strip():  # Ensure it's not an empty paragraph
                        try:
                            # Tách các phần như "A.", "B." và giữ nguyên chúng
                            preserved_parts = re.findall(r'^[A-Z]\.', para.text)
                            if preserved_parts:
                                cleaned_text = re.sub(r'^[A-Z]\.', '', para.text).strip()
                                translated_text_fragment = translator.translate(cleaned_text, src=source_language, dest=target_language).text
                                translated_paragraph = f"{preserved_parts[0]} {translated_text_fragment}"
                            else:
                                translated_paragraph = translator.translate(para.text, src=source_language, dest=target_language).text

                            para.text = translated_paragraph  # Update the paragraph with translated text
                            translated_text += translated_paragraph + "\n"  # Concatenate all translated text
                        except Exception as e:
                            print(f"Error during translation: {e}")
                            para.text = para.text  # Keep the original text in case of error
                            translated_text += para.text + "\n"  # Keep the original text in the concatenated result

                # Save the translated DOCX
                translated_docx_path = os.path.join(tempfile.gettempdir(), f"{original_filename}_translated.docx")
                doc.save(translated_docx_path)

                # Save the path for download
                request.session['translated_docx_path'] = translated_docx_path

                # Save the concatenated translated text for preview (if needed)
                request.session['translated_text'] = translated_text

            finally:
                # Clean up temporary files
                os.remove(pdf_path)

            return redirect('result_trans')
    else:
        form = PDFUploadForm()

    return render(request, 'translate/pdf_translate.html', {'form': form})



def convert_and_translate_docx_view(request):
    if request.method == 'POST':
        form = DOCXUploadForm(request.POST, request.FILES)
        if form.is_valid():
            docx_file = request.FILES['docx_file']
            original_filename = docx_file.name.rsplit('.', 1)[0]  # Original filename without extension
            source_language = form.cleaned_data['source_language']
            target_language = form.cleaned_data['target_language']

            # Save the DOCX file to a temporary location
            docx_path = os.path.join(tempfile.gettempdir(), f"{original_filename}.docx")
            with open(docx_path, 'wb') as temp_docx:
                for chunk in docx_file.chunks():
                    temp_docx.write(chunk)

            try:
                # Read the content from the DOCX file
                doc = Document(docx_path)

                # Translate the text using Google Translate API (if needed)
                translated_text = ""
                translator = Translator()

                # Translate and update the text in the DOCX
                for para in doc.paragraphs:
                    if para.text.strip():  # Ensure it's not an empty paragraph
                        try:
                            # Tách các phần như "A.", "B." và giữ nguyên chúng
                            preserved_parts = re.findall(r'^[A-Z]\.', para.text)
                            if preserved_parts:
                                cleaned_text = re.sub(r'^[A-Z]\.', '', para.text).strip()  # Remove "A." or "B." for translation
                                translated_text_fragment = translator.translate(cleaned_text, src=source_language, dest=target_language).text
                                translated_paragraph = f"{preserved_parts[0]} {translated_text_fragment}"
                            else:
                                translated_paragraph = translator.translate(para.text, src=source_language, dest=target_language).text

                            para.text = translated_paragraph  # Update the paragraph with translated text
                            translated_text += translated_paragraph + "\n"  # Concatenate all translated text
                        except Exception as e:
                            print(f"Error during translation: {e}")
                            para.text = para.text  # Keep the original text in case of error
                            translated_text += para.text + "\n"  # Keep the original text in the concatenated result

                # Save the translated DOCX
                translated_docx_path = os.path.join(tempfile.gettempdir(), f"{original_filename}_translated.docx")
                doc.save(translated_docx_path)

                # Save the path for download
                request.session['translated_docx_path'] = translated_docx_path

                # Save the concatenated translated text for preview (if needed)
                request.session['translated_text'] = translated_text

            finally:
                # Clean up the temporary file
                os.remove(docx_path)

            return redirect('result_trans')
    else:
        form = DOCXUploadForm()

    return render(request, 'translate/docx_translate.html', {'form': form})


def download_translated_docx_view(request):
    translated_docx_path = request.session.get('translated_docx_path')
    if not translated_docx_path or not os.path.exists(translated_docx_path):
        return HttpResponse("Không tìm thấy file để tải xuống. Vui lòng thử lại.")

    with open(translated_docx_path, 'rb') as docx_file:
        response = HttpResponse(
            docx_file.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(translated_docx_path)}"'
        return response


def result_view(request):
    translated_text = request.session.get('translated_text')
    if not translated_text:
        return HttpResponse("Không tìm thấy kết quả. Vui lòng thử lại.")

    return render(request, 'translate/result.html', {'docx_html': translated_text})

