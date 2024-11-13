import cv2
import numpy as np
import pytesseract
from docx import Document
from pdf2image import convert_from_path
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from process_img.utils import  noise_removal, enhance_text
import tempfile
import io


import tempfile



def process_uploaded_pdf(uploaded_pdf):
    doc = Document()  # Tạo file Word để lưu kết quả OCR

    # Tạo file tạm thời từ InMemoryUploadedFile
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
        temp_pdf.write(uploaded_pdf.read())  # Lưu nội dung file PDF vào file tạm thời
        temp_pdf_path = temp_pdf.name  # Đường dẫn của file tạm thời

    try:
        # Chuyển đổi PDF thành các ảnh với độ phân giải 300 DPI, mà không lưu ảnh vào thư mục
        pages = convert_from_path(temp_pdf_path, 300, fmt="png")

        for i, page in enumerate(pages):
            # Lưu ảnh của trang vào bộ nhớ (dưới dạng đối tượng bytes)
            img_bytes = io.BytesIO()
            page.save(img_bytes, format="PNG")
            img_bytes.seek(0)

            # Chuyển đổi ảnh từ bộ nhớ thành đối tượng OpenCV để xử lý
            file_bytes = np.asarray(bytearray(img_bytes.read()), dtype=np.uint8)
            img = cv2.imdecode(file_bytes, cv2.IMREAD_GRAYSCALE)

            # # Áp dụng giảm nhiễu
            # img = noise_removal(img)
            # Làm đậm nét chữ
            enhanced_img = enhance_text(img)

            # Áp dụng threshold để chuyển ảnh thành đen trắng
            _, binary_img = cv2.threshold(enhanced_img, 150, 255, cv2.THRESH_BINARY_INV)



            # Tìm các đường kẻ ngang và dọc để xác định bảng
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
            horizontal_lines = cv2.morphologyEx(binary_img, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
            vertical_lines = cv2.morphologyEx(binary_img, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
            table_area = cv2.add(horizontal_lines, vertical_lines)

            # Tìm contours để phân chia bảng
            contours, _ = cv2.findContours(table_area, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
            bounding_boxes = [cv2.boundingRect(c) for c in contours]



            if bounding_boxes:
                max_contour = max(bounding_boxes, key=lambda b: b[2] * b[3])
                bounding_boxes = [b for b in bounding_boxes if b != max_contour]
                img_text_only = enhanced_img.copy()
                for x, y, w, h in bounding_boxes:
                    cv2.rectangle(img_text_only, (x, y), (x + w, y + h), (0, 0, 0), -1)

                # OCR cho văn bản ngoài bảng
                text_from_image = pytesseract.image_to_string(img_text_only, lang='vie').strip()
                doc.add_paragraph(text_from_image)

                # OCR cho từng ô trong bảng
                table_data = []
                for x, y, w, h in bounding_boxes:
                    if w > 50 and h > 20:
                        cell_img = img[y:y + h, x:x + w]
                        text = pytesseract.image_to_string(cell_img, lang='vie').strip()
                        table_data.append((x, y, text))

                # Tạo bảng trong file Word
                if table_data:
                    rows = len(set([y for _, y, _ in table_data]))
                    cols = len(set([x for x, _, _ in table_data]))
                    table = doc.add_table(rows=rows, cols=cols)

                    sorted_table_data = sorted(table_data, key=lambda t: (t[1], t[0]))
                    row_idx, col_idx = 0, 0
                    prev_y = sorted_table_data[0][1]

                    for x, y, text in sorted_table_data:
                        if y != prev_y:
                            row_idx += 1
                            col_idx = 0
                            prev_y = y

                        if row_idx < rows and col_idx < cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = text
                            cell._element.get_or_add_tcPr().append(parse_xml(
                                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:left w:val="single" w:sz="4"/><w:right w:val="single" w:sz="4"/>'
                                r'<w:top w:val="single" w:sz="4"/><w:bottom w:val="single" w:sz="4"/>'
                                r'</w:tcBorders>'
                            ))
                            col_idx += 1

            else:
                # Nếu không có bảng, chỉ OCR cho toàn bộ ảnh
                text_from_image = pytesseract.image_to_string(binary_img, lang='vie').strip()
                doc.add_paragraph(text_from_image)

        downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
        output_file_path = os.path.join(downloads_folder, 'ocr_result.docx')
        doc.save(output_file_path)

        # Chuyển đổi file .docx thành HTML
        html_content = convert_docx_to_html(output_file_path)
        return html_content  # Trả về nội dung HTML để hiển thị

    finally:
        # Xóa file tạm sau khi xử lý xong
        os.remove(temp_pdf_path)



def convert_docx_to_html(doc_path):
    doc = Document(doc_path)
    html_content = "<div>"

    # Lặp qua từng đoạn văn bản trong file .docx
    for para in doc.paragraphs:
        html_content += f"<p>{para.text}</p>"

    # Lặp qua từng bảng và tạo cấu trúc HTML cho bảng
    for table in doc.tables:
        html_content += "<table border='1' style='border-collapse: collapse; width: 100%;'>"
        for row in table.rows:
            html_content += "<tr>"
            for cell in row.cells:
                html_content += f"<td>{cell.text}</td>"
            html_content += "</tr>"
        html_content += "</table>"

    html_content += "</div>"
    return html_content








# def process_uploaded_pdf(uploaded_pdf):
#     doc = Document()  # Tạo file Word để lưu kết quả OCR
#
#     # Tạo file tạm thời từ InMemoryUploadedFile
#     with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
#         temp_pdf.write(uploaded_pdf.read())  # Lưu nội dung file PDF vào file tạm thời
#         temp_pdf_path = temp_pdf.name  # Đường dẫn của file tạm thời
#
#     try:
#         # Chuyển đổi PDF thành các ảnh với độ phân giải 300 DPI, mà không lưu ảnh vào thư mục
#         pages = convert_from_path(temp_pdf_path, 300, fmt="png")
#
#         for i, page in enumerate(pages):
#             # Lưu ảnh của trang vào bộ nhớ (dưới dạng đối tượng bytes)
#             img_bytes = io.BytesIO()
#             page.save(img_bytes, format="PNG")
#             img_bytes.seek(0)
#
#             # Chuyển đổi ảnh từ bộ nhớ thành đối tượng OpenCV để xử lý
#             file_bytes = np.asarray(bytearray(img_bytes.read()), dtype=np.uint8)
#             img = cv2.imdecode(file_bytes, cv2.IMREAD_GRAYSCALE)
#
#             # Làm đậm nét chữ
#             enhanced_img = enhance_text(img)
#
#             # Áp dụng threshold để chuyển ảnh thành đen trắng
#             _, binary_img = cv2.threshold(enhanced_img, 150, 255, cv2.THRESH_BINARY_INV)
#
#             # Tìm các đường kẻ ngang và dọc để xác định bảng
#             horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
#             vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
#             horizontal_lines = cv2.morphologyEx(binary_img, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)
#             vertical_lines = cv2.morphologyEx(binary_img, cv2.MORPH_OPEN, vertical_kernel, iterations=2)
#             table_area = cv2.add(horizontal_lines, vertical_lines)
#
#             # Tìm contours để phân chia bảng
#             contours, _ = cv2.findContours(table_area, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
#             bounding_boxes = [cv2.boundingRect(c) for c in contours]
#
#             sections = []
#             prev_y = 0  # Lưu trữ vị trí Y của phần trước để xác định các vùng văn bản và bảng
#             index = 1  # Bắt đầu từ 1 cho phần đầu tiên (văn bản trước bảng)
#
#             # Thêm phần văn bản đầu tiên (trước bảng)
#             if bounding_boxes:
#                 first_box = bounding_boxes[0]
#                 sections.append((index, (0, 0, first_box[0], img.shape[0])))  # Văn bản trước bảng
#
#             # Thêm các bảng và các phần sau bảng
#             index = 2  # Bắt đầu đánh index cho bảng
#             for x, y, w, h in bounding_boxes:
#                 sections.append((index, (x, y, w, h)))  # Bảng (index 2, 3,...)
#                 index += 1
#
#             # Thêm phần văn bản còn lại sau bảng (nếu có)
#             if bounding_boxes:
#                 last_box = bounding_boxes[-1]
#                 sections.append((index, (last_box[0] + last_box[2], 0, img.shape[1], img.shape[0])))
#
#             # Thực hiện OCR cho từng phần
#             for index, region in sections:
#                 x, y, w, h = region
#                 if index == 1:  # OCR cho phần văn bản trước bảng
#                     text_from_image = pytesseract.image_to_string(enhanced_img[y:y+h, x:x+w], lang='vie').strip()
#                     doc.add_paragraph(text_from_image)
#                 elif index > 1:  # OCR cho bảng
#                     table_data = []
#                     for bx, by, bw, bh in bounding_boxes:
#                         if bw > 50 and bh > 20:
#                             cell_img = enhanced_img[by:by+bh, bx:bx+bw]
#                             text = pytesseract.image_to_string(cell_img, lang='vie').strip()
#                             table_data.append((bx, by, text))
#
#                     if table_data:
#                         rows = len(set([y for _, y, _ in table_data]))
#                         cols = len(set([x for x, _, _ in table_data]))
#                         table = doc.add_table(rows=rows, cols=cols)
#
#                         sorted_table_data = sorted(table_data, key=lambda t: (t[1], t[0]))
#                         row_idx, col_idx = 0, 0
#                         prev_y = sorted_table_data[0][1]
#
#                         for bx, by, text in sorted_table_data:
#                             if by != prev_y:
#                                 row_idx += 1
#                                 col_idx = 0
#                                 prev_y = by
#
#                             if row_idx < rows and col_idx < cols:
#                                 cell = table.cell(row_idx, col_idx)
#                                 cell.text = text
#                                 col_idx += 1
#
#         downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
#         output_file_path = os.path.join(downloads_folder, 'ocr_result.docx')
#         doc.save(output_file_path)
#
#         # Chuyển đổi file .docx thành HTML
#         html_content = convert_docx_to_html(output_file_path)
#         return html_content  # Trả về nội dung HTML để hiển thị
#
#     finally:
#         # Xóa file tạm sau khi xử lý xong
#         os.remove(temp_pdf_path)
#
#
#
#
# def convert_docx_to_html(doc_path):
#     doc = Document(doc_path)
#     html_content = "<div>"
#
#     # Lặp qua từng đoạn văn bản trong file .docx
#     for para in doc.paragraphs:
#         html_content += f"<p>{para.text}</p>"
#
#     # Lặp qua từng bảng và tạo cấu trúc HTML cho bảng
#     for table in doc.tables:
#         html_content += "<table border='1' style='border-collapse: collapse; width: 100%;'>"
#         for row in table.rows:
#             html_content += "<tr>"
#             for cell in row.cells:
#                 html_content += f"<td>{cell.text}</td>"
#             html_content += "</tr>"
#         html_content += "</table>"
#
#     html_content += "</div>"
#     return html_content


